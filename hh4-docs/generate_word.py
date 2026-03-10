"""
generate_word.py
Generates HH4_Platform_Solution_Architecture.docx
Run: python generate_word.py
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from copy import deepcopy

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x1B, 0x2A, 0x4A)   # #1B2A4A
GOLD   = RGBColor(0xC9, 0xA8, 0x4C)   # #C9A84C
BLUE2  = RGBColor(0x2E, 0x50, 0x90)   # #2E5090
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
BLACK  = RGBColor(0x00, 0x00, 0x00)
LIGHT  = RGBColor(0xF0, 0xF4, 0xF8)   # #F0F4F8 alternating row

NAVY_HEX  = "1B2A4A"
GOLD_HEX  = "C9A84C"
LIGHT_HEX = "F0F4F8"
WHITE_HEX = "FFFFFF"
BLUE2_HEX = "2E5090"


# ── XML helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Fill a table cell background with a solid colour (hex, no #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    # remove existing shd if any
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(shd)


def set_cell_borders(cell, color_hex="AAAAAA"):
    """Apply thin borders to all sides of a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color_hex)
        tcBorders.append(el)
    for old in tcPr.findall(qn("w:tcBorders")):
        tcPr.remove(old)
    tcPr.append(tcBorders)


def add_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type("page"))


def docx_break_type(kind="page"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), kind)
    return br


def insert_page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run._r.append(docx_break_type("page"))
    return para


# ── Typography helpers ────────────────────────────────────────────────────────

def set_run_font(run, size_pt, bold=False, italic=False, color=None, font_name="Calibri"):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    """Add a styled heading paragraph (not using built-in Heading styles)."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run(text)
    if level == 1:
        set_run_font(run, 20, bold=True, color=NAVY)
    elif level == 2:
        set_run_font(run, 16, bold=True, color=NAVY)
    elif level == 3:
        set_run_font(run, 13, bold=True, color=BLUE2)
    return para


def add_body(doc, text, italic=False, space_after=6):
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(space_after)
    run = para.add_run(text)
    set_run_font(run, 11, italic=italic, color=BLACK)
    return para


def add_bullet(doc, text, symbol="•"):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(0.8)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(f"{symbol}  {text}")
    set_run_font(run, 11, color=BLACK)
    return para


# ── Table helpers ─────────────────────────────────────────────────────────────

def make_table(doc, headers, rows, col_widths=None, border_color="AAAAAA"):
    """
    Create a styled table.
    headers : list of str  (header row)
    rows    : list of list of str (data rows, alternating shading)
    col_widths : list of Cm values or None (auto)
    """
    col_count = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=col_count)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        set_cell_bg(cell, NAVY_HEX)
        set_cell_borders(cell, border_color)
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(h)
        set_run_font(run, 10, bold=True, color=WHITE)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        bg = LIGHT_HEX if r_idx % 2 == 0 else WHITE_HEX
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            cell = cells[c_idx]
            set_cell_bg(cell, bg)
            set_cell_borders(cell, border_color)
            para = cell.paragraphs[0]
            run = para.add_run(str(val))
            set_run_font(run, 10, color=BLACK)

    # Column widths
    if col_widths:
        for r in table.rows:
            for c_idx, w in enumerate(col_widths):
                if c_idx < len(r.cells):
                    r.cells[c_idx].width = w

    doc.add_paragraph()  # spacer
    return table


# ── Header / Footer ───────────────────────────────────────────────────────────

def add_header_footer(doc):
    section = doc.sections[0]

    # Header
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = hp.add_run("Splendora HH4 Digital Platform — Confidential")
    set_run_font(run, 9, italic=True, color=NAVY)

    # Footer (page number via field)
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run()
    set_run_font(run, 9, color=NAVY)
    # Insert PAGE field
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instrText = OxmlElement("w:instrText")
    instrText.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


# ── Document setup ────────────────────────────────────────────────────────────

def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)
    # Default paragraph font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


# ══════════════════════════════════════════════════════════════════════════════
#  COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════

def add_cover(doc):
    # top spacer
    for _ in range(6):
        doc.add_paragraph()

    # Main title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SPLENDORA HH4 DIGITAL PLATFORM")
    set_run_font(r, 32, bold=True, color=NAVY)

    # Gold divider
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("━" * 40)
    set_run_font(r2, 14, color=GOLD)

    # Subtitle
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Solution Architecture & Investment Proposal")
    set_run_font(r3, 18, italic=True, color=NAVY)

    doc.add_paragraph()

    # Investment value
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r4 = p4.add_run("Investment Value: >$10,000,000 USD")
    set_run_font(r4, 14, bold=True, color=GOLD)

    for _ in range(4):
        doc.add_paragraph()

    # Metadata line
    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r5 = p5.add_run("March 2026  |  Classification: Strictly Confidential  |  Version: 1.0")
    set_run_font(r5, 11, italic=True, color=NAVY)

    insert_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  TABLE OF CONTENTS
# ══════════════════════════════════════════════════════════════════════════════

def add_toc(doc):
    add_heading(doc, "TABLE OF CONTENTS", 1)
    toc_entries = [
        ("1",  "Executive Summary",                   "3"),
        ("2",  "Why This Investment?",                "4"),
        ("3",  "Solution Architecture",               "6"),
        ("4",  "Functional Module Breakdown",         "12"),
        ("5",  "User Roles & Permission Matrix",      "14"),
        ("6",  "KPI Framework",                       "15"),
        ("7",  "Business Case & Financial Model",     "18"),
        ("8",  "Phased Roadmap",                      "22"),
        ("9",  "RFP & Vendor Selection",              "24"),
        ("10", "Risk Mitigation",                     "26"),
        ("11", "Competitive Differentiation",         "27"),
    ]
    for num, title, page in toc_entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        tab_stops = p.paragraph_format.tab_stops
        tab_stops.add_tab_stop(Cm(14), WD_ALIGN_PARAGRAPH.RIGHT)
        run = p.add_run(f"Section {num}  —  {title}")
        set_run_font(run, 11, color=BLACK)
        run2 = p.add_run(f"\t{page}")
        set_run_font(run2, 11, color=NAVY)

    insert_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 – EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════

def add_section1(doc):
    add_heading(doc, "SECTION 1: EXECUTIVE SUMMARY", 1)
    add_body(doc,
        "The HH4 Digital Platform represents a transformational investment in Splendora's "
        "competitive positioning, delivering an integrated ecosystem of AI-driven real estate "
        "experiences that directly accelerates revenue, elevates brand equity, and creates a "
        "defensible moat in the Vietnam premium residential market.")

    headers = ["Hạng mục", "Nội dung"]
    rows = [
        ["Mục tiêu chiến lược",
         "Xây dựng hệ sinh thái số tích hợp AI cho phân khúc BĐS cao cấp HH4 Splendora"],
        ["Quy mô đầu tư",
         ">$10,000,000 USD — bao gồm platform, AI engine, content, vận hành 3 năm"],
        ["Sản phẩm cốt lõi",
         "IVA Metaverse Showroom + Concierge 360° Resident App"],
        ["Đối tượng phục vụ",
         "Khách mua BĐS cao cấp, cư dân hiện hữu, đội ngũ sales & quản lý vận hành"],
        ["Kiến trúc",
         "Cloud-native, microservices, AI/ML pipeline, zero-trust security"],
        ["Timeline",
         "18 tháng triển khai, Go-live Q3 2026, Full ROI dự kiến Year 3"],
    ]
    make_table(doc, headers, rows, col_widths=[Cm(5), Cm(12)])


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 – WHY THIS INVESTMENT?
# ══════════════════════════════════════════════════════════════════════════════

def add_section2(doc):
    add_heading(doc, "SECTION 2: WHY THIS INVESTMENT?", 1)

    # 2.1
    add_heading(doc, "2.1  Market Problems", 2)
    problems = [
        "Showroom truyền thống không thể phục vụ khách hàng quốc tế và Việt kiều từ xa",
        "Quy trình bán hàng BĐS cao cấp còn phân mảnh, thiếu dữ liệu hành vi khách hàng",
        "Trải nghiệm cư dân sau bàn giao thiếu kết nối số, dẫn đến churn cao và NPS thấp",
        "Đối thủ cạnh tranh quốc tế đang tích hợp AI và metaverse vào quy trình bán hàng",
        "Thiếu nền tảng dữ liệu thống nhất để ra quyết định đầu tư và marketing chính xác",
    ]
    for p in problems:
        add_bullet(doc, p, "❌")

    # 2.2
    add_heading(doc, "2.2  Market Opportunities", 2)
    headers = ["Đòn bẩy", "Giá trị tạo ra"]
    rows = [
        ["AI-powered virtual showroom",
         "Tăng qualified lead rate 3×, giảm cost per lead 40%"],
        ["Metaverse & 3D tour",
         "Tiếp cận khách quốc tế không giới hạn địa lý, tăng reach 5×"],
        ["Concierge 360° app",
         "Tăng resident retention 25%, tạo recurring revenue từ dịch vụ"],
        ["Behavioral AI scoring",
         "Tăng conversion rate sales từ 2% lên 6%, rút ngắn sales cycle 30%"],
        ["Data monetisation",
         "Mở ra B2B data partnerships, partner GMV >$2M năm 3"],
    ]
    make_table(doc, headers, rows, col_widths=[Cm(7), Cm(10)])

    # 2.3
    add_heading(doc, "2.3  ROI Projection", 2)
    headers = ["Metric", "Giá trị"]
    rows = [
        ["Total GTV (Gross Transaction Value) — portfolio HH4",    "$250,000,000 USD"],
        ["Baseline conversion rate (hiện tại)",                    "2.0%"],
        ["Target conversion rate (với HH4 Platform)",              "6.0%"],
        ["Delta Revenue (incremental sales)",                      "+$10,000,000 USD"],
        ["Platform total cost (3 năm)",                            "$10,500,000 USD"],
        ["Concierge recurring revenue (3 năm)",                    "$3,600,000 USD"],
        ["Net ROI (Year 3)",                                       "+29% trên tổng đầu tư"],
    ]
    make_table(doc, headers, rows, col_widths=[Cm(9), Cm(8)])

    insert_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 – SOLUTION ARCHITECTURE
# ══════════════════════════════════════════════════════════════════════════════

def add_section3(doc):
    add_heading(doc, "SECTION 3: SOLUTION ARCHITECTURE", 1)

    # 3.1
    add_heading(doc, "3.1  Ecosystem Overview", 2)
    add_body(doc,
        "The HH4 Platform is built as a cloud-native, AI-first ecosystem composed of two "
        "flagship product lines — the IVA Metaverse Showroom targeting pre-purchase journey "
        "optimisation, and the Concierge 360° Resident App governing post-handover lifecycle. "
        "Both share a unified data lake, AI/ML pipeline, and identity layer.")

    ascii_art = (
        "┌─────────────────────────────────────────────────────────────┐\n"
        "│                  HH4 DIGITAL PLATFORM                      │\n"
        "│                                                             │\n"
        "│   ┌──────────────────┐       ┌───────────────────────┐     │\n"
        "│   │  IVA METAVERSE   │       │   CONCIERGE 360°       │     │\n"
        "│   │    SHOWROOM      │       │   RESIDENT APP         │     │\n"
        "│   │  (Pre-purchase)  │       │   (Post-handover)      │     │\n"
        "│   └────────┬─────────┘       └──────────┬────────────┘     │\n"
        "│            │                            │                  │\n"
        "│   ─────────┴────────────────────────────┴─────────────     │\n"
        "│              UNIFIED AI / DATA PLATFORM                    │\n"
        "│   ─────────────────────────────────────────────────────    │\n"
        "│    Identity  │  Analytics  │  AI Engine  │  CRM/ERP       │\n"
        "└─────────────────────────────────────────────────────────────┘"
    )
    p = doc.add_paragraph()
    r = p.add_run(ascii_art)
    r.font.name = "Courier New"
    r.font.size = Pt(8)
    r.font.color.rgb = NAVY

    # 3.2 — IVA Metaverse Showroom
    add_heading(doc, "3.2  MODULE 1 — IVA Metaverse Showroom", 2)
    add_body(doc,
        "An AI-powered, photorealistic 3D virtual showroom enabling prospects worldwide "
        "to explore HH4 units, interact with an intelligent virtual assistant, and progress "
        "through a guided discovery journey — fully tracked and scored by the behavioral AI engine.")

    add_heading(doc, "Customer Journey Flow", 3)
    headers = ["Stage", "Tâm lý KH", "Điểm chạm", "Tính năng", "KPI"]
    rows = [
        ["1 · Awareness",
         "Tò mò, khám phá",
         "QR / Landing page / Social ad",
         "3D lobby entry, ambient tour",
         "Entry Rate ≥ 35%"],
        ["2 · Exploration",
         "Đánh giá, so sánh",
         "Virtual floor plan, zone hotspot",
         "AI tour guide, zone highlights",
         "Zone Completion ≥ 60%"],
        ["3 · Engagement",
         "Quan tâm cụ thể",
         "Hotspot interaction, live chat",
         "AI chat, document download",
         "Hotspot Opens ≥ 4/session"],
        ["4 · Qualification",
         "Cân nhắc nghiêm túc",
         "Behavioral score alert",
         "AI scoring, sales notification",
         "Qualified Lead Rate ≥ 15%"],
        ["5 · Appointment",
         "Sẵn sàng gặp sales",
         "Booking widget in showroom",
         "Calendar sync, reminder",
         "Tour-booking Rate ≥ 8%"],
        ["6 · Decision",
         "Chốt deal",
         "VIP virtual consultation room",
         "1-on-1 video, e-signature",
         "Conversion to Deposit ≥ 3%"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(2.5), Cm(3), Cm(3), Cm(4), Cm(4)])

    add_heading(doc, "AI Behavioral Scoring Engine", 3)
    headers = ["Signal", "Weight", "Threshold → Action"]
    rows = [
        ["Time-on-tour (minutes)",          "20%", ">8 min → qualify"],
        ["Hotspot interactions",            "15%", ">5 → notify sales"],
        ["Floor-plan downloads",            "15%", "Any → hot lead"],
        ["Pricing page visits",             "20%", ">2 → high intent"],
        ["Return visits within 7 days",     "15%", "≥2 → priority lead"],
        ["Live chat initiated",             "10%", "Any → immediate follow-up"],
        ["Booking widget interaction",      "5%",  "Any → urgent follow-up"],
    ]
    make_table(doc, headers, rows, col_widths=[Cm(6), Cm(3), Cm(8)])

    # 3.3 — Concierge 360°
    add_heading(doc, "3.3  MODULE 2 — Concierge 360°", 2)
    add_body(doc,
        "A super-app for HH4 residents delivering AI-powered service booking, smart home "
        "control, community engagement, loyalty rewards, and predictive maintenance — "
        "transforming the resident relationship from transactional to lifestyle partnership.")

    add_heading(doc, "Resident Journey", 3)
    headers = ["Stage", "Kỳ vọng", "Hành động", "Tính năng", "KPI"]
    rows = [
        ["1 · Onboarding",
         "Dễ dàng bắt đầu",
         "Đăng ký, xác thực căn hộ",
         "eKYC, profile setup wizard",
         "Activation Rate ≥ 70%"],
        ["2 · Discovery",
         "Khám phá tiện ích",
         "Browse services, AI recommendation",
         "Personalised feed, service catalogue",
         "First Interaction ≤ 48h"],
        ["3 · Usage",
         "Tiện lợi, đáng tin cậy",
         "Đặt dịch vụ, thanh toán",
         "One-tap booking, digital wallet",
         "Booking Frequency ≥ 3×/month"],
        ["4 · Smart Living",
         "Kiểm soát thông minh",
         "Điều khiển smarthome, tiết kiệm NL",
         "IoT dashboard, energy analytics",
         "Smart Home Usage ≥ 50%"],
        ["5 · Community",
         "Kết nối hàng xóm",
         "Tham gia sự kiện, diễn đàn",
         "Event RSVP, community board",
         "Event Participation ≥ 25%"],
        ["6 · Advocacy",
         "Giới thiệu bạn bè",
         "Chia sẻ, referral",
         "Loyalty points, referral engine",
         "Advocacy Conversion ≥ 5%"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(2.5), Cm(3), Cm(3.5), Cm(4), Cm(3.5)])

    add_heading(doc, "AI Concierge Engine", 3)
    add_body(doc,
        "The AI Concierge Engine uses a multi-modal LLM backbone (fine-tuned on Vietnamese "
        "real-estate and property-management domain data) to deliver: proactive service "
        "recommendations, predictive maintenance scheduling, natural-language service requests "
        "via voice/chat, real-time sentiment analysis from feedback loops, and cross-sell "
        "partner offers matched to resident lifestyle profiles.")

    # 3.4 — Platform Infrastructure
    add_heading(doc, "3.4  Platform Infrastructure", 2)
    headers = ["Layer", "Technology Stack", "Notes"]
    rows = [
        ["Frontend — Web",
         "React 18, TypeScript, WebGL/Three.js, WebXR",
         "PWA, SSR, <2s load time"],
        ["Frontend — Mobile",
         "React Native (iOS/Android), Expo",
         "Offline-first, biometric auth"],
        ["Frontend — VR/AR",
         "Unity 2023 (WebGL build), A-Frame",
         "Browser-native, no plugin"],
        ["Backend — API",
         "Node.js / FastAPI microservices, GraphQL",
         "OpenAPI 3.1 spec"],
        ["Backend — Orchestration",
         "Kubernetes (GKE), Istio service mesh",
         "Auto-scaling, zero-downtime"],
        ["Data — Storage",
         "PostgreSQL, Redis, MongoDB Atlas",
         "Multi-region replication"],
        ["Data — Lake",
         "BigQuery, Apache Kafka, dbt",
         "Real-time + batch pipelines"],
        ["AI/ML",
         "Vertex AI, TensorFlow, LangChain, GPT-4o",
         "Fine-tuned VN domain models"],
        ["DevOps",
         "GitHub Actions, ArgoCD, Terraform",
         "GitOps, IaC"],
    ]
    make_table(doc, headers, rows, col_widths=[Cm(4), Cm(7), Cm(5)])

    # 3.5 — Security Architecture
    add_heading(doc, "3.5  Security Architecture", 2)
    headers = ["Layer", "Control", "Standard / Tool"]
    rows = [
        ["Network",    "WAF, DDoS protection, private VPC",          "Cloudflare, GCP VPC"],
        ["Network",    "Zero-trust network access",                  "BeyondCorp / Tailscale"],
        ["Application","OAuth 2.0 + PKCE, MFA, biometric",          "Auth0 / Firebase Auth"],
        ["Application","API rate limiting, input validation",        "Kong Gateway"],
        ["Application","SAST/DAST in CI/CD pipeline",               "Snyk, OWASP ZAP"],
        ["Data",       "AES-256 encryption at rest",                 "GCP KMS"],
        ["Data",       "TLS 1.3 in transit",                        "Let's Encrypt / cert-manager"],
        ["Data",       "PDPA / GDPR compliance, data minimisation",  "OneTrust, DPO oversight"],
        ["Monitoring", "SIEM, anomaly detection, SOC alerts",        "Google Chronicle, PagerDuty"],
    ]
    make_table(doc, headers, rows, col_widths=[Cm(3), Cm(8), Cm(5)])

    insert_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 – FUNCTIONAL MODULE BREAKDOWN
# ══════════════════════════════════════════════════════════════════════════════

def add_section4(doc):
    add_heading(doc, "SECTION 4: FUNCTIONAL MODULE BREAKDOWN", 1)
    add_body(doc,
        "The matrix below maps all 20 core platform functions to the five product modules. "
        "✅ = fully delivered within scope; ❌ = out of scope for this phase.")

    headers = [
        "Function",
        "Marketing",
        "Experience",
        "Sales Enablement",
        "Concierge & Service",
        "Admin & Ops",
    ]
    rows = [
        ["3D Virtual Showroom",                 "✅", "✅", "✅", "❌", "❌"],
        ["AI Virtual Tour Guide",               "✅", "✅", "✅", "❌", "❌"],
        ["Behavioral Lead Scoring",             "✅", "❌", "✅", "❌", "✅"],
        ["CRM Integration",                     "❌", "❌", "✅", "❌", "✅"],
        ["Live Video Consultation",             "❌", "✅", "✅", "❌", "❌"],
        ["E-Signature & Contract",              "❌", "❌", "✅", "❌", "✅"],
        ["Smart Home Dashboard",                "❌", "❌", "❌", "✅", "❌"],
        ["Service Booking Engine",              "❌", "❌", "❌", "✅", "✅"],
        ["Digital Wallet / Payments",           "❌", "❌", "❌", "✅", "✅"],
        ["Loyalty & Rewards",                   "✅", "❌", "❌", "✅", "✅"],
        ["Community & Social Feed",             "❌", "✅", "❌", "✅", "❌"],
        ["Event Management",                    "✅", "✅", "❌", "✅", "✅"],
        ["AI Chatbot (Vietnamese NLP)",         "✅", "✅", "✅", "✅", "❌"],
        ["Push Notifications",                  "✅", "❌", "✅", "✅", "✅"],
        ["Analytics Dashboard",                 "✅", "❌", "✅", "✅", "✅"],
        ["Partner Marketplace",                 "❌", "❌", "❌", "✅", "✅"],
        ["Predictive Maintenance Alerts",       "❌", "❌", "❌", "✅", "✅"],
        ["Multi-language (EN / VI / ZH)",       "✅", "✅", "✅", "✅", "❌"],
        ["Role-Based Access Control",           "❌", "❌", "❌", "❌", "✅"],
        ["Audit Log & Compliance Reporting",    "❌", "❌", "❌", "❌", "✅"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(5.5), Cm(2.5), Cm(2.5), Cm(3.5), Cm(3.5), Cm(2.5)])

    insert_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 – USER ROLES & PERMISSION MATRIX
# ══════════════════════════════════════════════════════════════════════════════

def add_section5(doc):
    add_heading(doc, "SECTION 5: USER ROLES & PERMISSION MATRIX", 1)
    headers = [
        "Permission",
        "Guest",
        "Registered",
        "VIP Member",
        "Sales / Mktg",
        "Admin",
    ]
    rows = [
        ["Browse 3D Showroom",          "✅", "✅", "✅", "✅", "✅"],
        ["Download Floor Plans",        "❌", "✅", "✅", "✅", "✅"],
        ["Request Pricing Sheet",       "❌", "✅", "✅", "✅", "✅"],
        ["Book Physical Tour",          "❌", "✅", "✅", "✅", "✅"],
        ["VIP Exclusive Areas",         "❌", "❌", "✅", "✅", "✅"],
        ["Access Lead Dashboard",       "❌", "❌", "❌", "✅", "✅"],
        ["Manage Listings",             "❌", "❌", "❌", "✅", "✅"],
        ["View Analytics",              "❌", "❌", "❌", "✅", "✅"],
        ["Service Booking (Concierge)", "❌", "✅", "✅", "❌", "✅"],
        ["Smart Home Control",          "❌", "✅", "✅", "❌", "✅"],
        ["Community Access",            "❌", "✅", "✅", "❌", "✅"],
        ["Loyalty Redemption",          "❌", "✅", "✅", "❌", "✅"],
        ["Partner Marketplace",         "❌", "✅", "✅", "❌", "✅"],
        ["User Management",             "❌", "❌", "❌", "❌", "✅"],
        ["System Configuration",        "❌", "❌", "❌", "❌", "✅"],
        ["Audit Log Access",            "❌", "❌", "❌", "❌", "✅"],
        ["Export Data",                 "❌", "❌", "❌", "✅", "✅"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(5), Cm(2), Cm(2.5), Cm(2.5), Cm(2.8), Cm(2)])

    insert_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 – KPI FRAMEWORK
# ══════════════════════════════════════════════════════════════════════════════

def add_section6(doc):
    add_heading(doc, "SECTION 6: KPI FRAMEWORK", 1)
    kpi_headers = ["Stage", "KPI Name", "Definition", "Target"]
    col_w = [Cm(3), Cm(4.5), Cm(7), Cm(3)]

    # 6.1 — IVA Funnel KPIs
    add_heading(doc, "6.1  KPI Funnel — IVA Metaverse Showroom", 2)
    iva_rows = [
        ["Awareness",     "Entry Rate",
         "% of ad/landing page visitors who enter the virtual showroom",
         "≥ 35%"],
        ["Awareness",     "Start-tour Rate",
         "% of entrants who begin the guided AI tour",
         "≥ 60%"],
        ["Exploration",   "Zone Completion Rate",
         "% of tours completing ≥ 3 distinct zones",
         "≥ 60%"],
        ["Exploration",   "Hotspot Opens",
         "Avg. number of hotspot interactions per session",
         "≥ 4"],
        ["Qualification", "Qualified Lead Rate",
         "% of sessions scoring above AI threshold",
         "≥ 15%"],
        ["Appointment",   "Tour-booking Rate",
         "% of qualified leads booking a physical or virtual appointment",
         "≥ 8%"],
        ["Appointment",   "Show-up Rate",
         "% of booked appointments that are attended",
         "≥ 75%"],
        ["Appointment",   "Appointment Set Rate",
         "% of sessions resulting in confirmed next step",
         "≥ 10%"],
        ["Decision",      "Shortlist Rate",
         "% of appointments resulting in unit shortlisting",
         "≥ 30%"],
        ["Decision",      "Reserve Intent",
         "% of shortlisted prospects expressing reservation intent",
         "≥ 20%"],
        ["Decision",      "Conversion to Deposit",
         "% of showroom entrants ultimately placing a deposit",
         "≥ 3%"],
    ]
    make_table(doc, kpi_headers, iva_rows, col_widths=col_w)

    # 6.2 — Concierge Funnel KPIs
    add_heading(doc, "6.2  KPI Funnel — Concierge 360°", 2)
    con_rows = [
        ["Onboarding",  "App Awareness Rate",
         "% of residents aware of the Concierge app at handover", "≥ 90%"],
        ["Onboarding",  "Account Activation Rate",
         "% of eligible residents completing account setup", "≥ 70%"],
        ["Onboarding",  "Profile Completion Rate",
         "% of activated accounts completing full profile", "≥ 65%"],
        ["Discovery",   "AI First Interaction Rate",
         "% of activated users who interact with AI Concierge within 48h", "≥ 60%"],
        ["Usage",       "Monthly Active Resident Rate (MAR)",
         "% of residents active in the app in a calendar month", "≥ 55%"],
        ["Usage",       "Service Booking Frequency",
         "Avg. number of service bookings per resident per month", "≥ 3×"],
        ["Usage",       "Payment Completion Rate",
         "% of initiated service bookings completed with payment", "≥ 85%"],
        ["Smart Living","Smart Home Usage Rate",
         "% of residents using IoT/smart home features monthly", "≥ 50%"],
        ["Smart Living","AI Recommendation Rate",
         "% of service bookings initiated by AI recommendation", "≥ 30%"],
        ["Community",   "Event Participation Rate",
         "% of residents RSVPing to ≥ 1 community event per quarter", "≥ 25%"],
        ["Community",   "Community Engagement Rate",
         "% of residents posting/commenting on community board monthly", "≥ 20%"],
        ["Loyalty",     "Loyalty Point Earn Rate",
         "% of monthly actives earning loyalty points", "≥ 70%"],
        ["Loyalty",     "Redemption Rate",
         "% of earned points redeemed within 90 days", "≥ 40%"],
        ["Loyalty",     "Cross-service Adoption Rate",
         "% of residents using ≥ 3 distinct service categories", "≥ 35%"],
        ["Partner",     "Partner GMV",
         "Gross merchandise value transacted via partner marketplace", "≥ $2M Y3"],
        ["Retention",   "Resident Retention Rate",
         "% of residents renewing / not selling within 24 months", "≥ 85%"],
        ["Retention",   "Churn Rate",
         "Monthly app churn (deactivations / active base)", "≤ 3%"],
        ["Satisfaction","NPS",
         "Net Promoter Score from quarterly resident survey", "≥ 60"],
        ["Advocacy",    "Repeat Service Rate",
         "% of residents booking same service category ≥ 2× in a month", "≥ 50%"],
        ["Advocacy",    "Advocacy Conversion Rate",
         "% of residents who refer ≥ 1 new prospect (sales or services)", "≥ 5%"],
        ["AI Quality",  "AI Satisfaction Score",
         "Avg. user rating of AI Concierge interactions (1–5)", "≥ 4.2"],
    ]
    make_table(doc, kpi_headers, con_rows, col_widths=col_w)

    insert_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 – BUSINESS CASE & FINANCIAL MODEL
# ══════════════════════════════════════════════════════════════════════════════

def add_section7(doc):
    add_heading(doc, "SECTION 7: BUSINESS CASE & FINANCIAL MODEL", 1)

    # 7.1
    add_heading(doc, "7.1  Investment Summary", 2)
    headers = ["Component", "Cost (USD)", "Timeline", "Owner", "Notes"]
    rows = [
        ["Platform Development (IVA + Concierge)", "$5,200,000",
         "M1–M12", "Technology Partner", "Fixed-price contract"],
        ["3D Content & Metaverse Assets",          "$1,800,000",
         "M3–M10", "Creative Agency",    "Photorealistic renders"],
        ["AI/ML Model Training & Fine-tuning",     "$800,000",
         "M6–M15", "AI Vendor",          "Vietnamese NLP"],
        ["Infrastructure (3-year cloud)",          "$1,200,000",
         "Ongoing", "Internal DevOps",   "GCP enterprise"],
        ["Operations & Support (Year 1–3)",        "$1,000,000",
         "Ongoing", "Splendora IT",      "SLA-backed support"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(5), Cm(3), Cm(2.5), Cm(3), Cm(4)])

    # 7.2
    add_heading(doc, "7.2  Revenue Model", 2)
    add_heading(doc, "Stream 1: Sales Acceleration Revenue", 3)
    headers = ["Metric", "Value"]
    rows = [
        ["Incremental units sold per year (vs baseline)", "+12 units"],
        ["Average unit price",                            "$800,000"],
        ["Gross incremental revenue (Year 1)",            "$9,600,000"],
        ["Commission / margin contribution (5%)",         "$480,000"],
    ]
    make_table(doc, headers, rows, col_widths=[Cm(9), Cm(8)])

    add_heading(doc, "Stream 2: Concierge Recurring Revenue", 3)
    headers = ["Revenue Item", "Year 1", "Year 2", "Year 3"]
    rows = [
        ["Subscription fee (residents)",   "$180,000", "$360,000",   "$540,000"],
        ["Partner commission (GMV 10%)",   "$80,000",  "$300,000",   "$600,000"],
        ["Premium service upsell",         "$60,000",  "$200,000",   "$400,000"],
        ["Advertising / sponsorship",      "$20,000",  "$100,000",   "$250,000"],
        ["Total",                          "$340,000", "$960,000",   "$1,790,000"],
    ]
    make_table(doc, headers, rows, col_widths=[Cm(6), Cm(3), Cm(3), Cm(3)])

    add_heading(doc, "Stream 3: Brand Premium & Land Value Uplift", 3)
    headers = ["Metric", "Estimate"]
    rows = [
        ["Price premium on HH4 units with digital showroom", "+3–5%"],
        ["Land / project valuation uplift",                  "+$15,000,000"],
        ["Future development cost reduction (data-driven)",  "-8% per sqm"],
    ]
    make_table(doc, headers, rows, col_widths=[Cm(9), Cm(8)])

    # 7.3
    add_heading(doc, "7.3  5-Year Financial Projection", 2)
    headers = ["Line Item", "Year 1", "Year 2", "Year 3", "Year 4", "Year 5"]
    rows = [
        ["Platform Investment (capex)",
         "($7,800K)", "($1,200K)", "($1,000K)", "($800K)", "($800K)"],
        ["Sales Delta Revenue",
         "$0", "$4,800K", "$9,600K", "$11,200K", "$12,800K"],
        ["Concierge Recurring Revenue",
         "$340K", "$960K", "$1,790K", "$2,400K", "$3,100K"],
        ["Brand Premium / Valuation",
         "$0", "$5,000K", "$10,000K", "$10,000K", "$10,000K"],
        ["Net Cash Flow",
         "($7,460K)", "$9,560K", "$20,390K", "$22,800K", "$25,100K"],
        ["Cumulative Cash Flow",
         "($7,460K)", "$2,100K", "$22,490K", "$45,290K", "$70,390K"],
        ["ROI (cumulative)",
         "-100%", "+28%", "+207%", "+476%", "+800%"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(4.5), Cm(2.3), Cm(2.3), Cm(2.3), Cm(2.3), Cm(2.3)])

    # 7.4
    add_heading(doc, "7.4  Cost Breakdown by Phase", 2)
    headers = ["Phase", "Line Item", "Cost (USD)"]
    rows = [
        ["Phase 1: Foundation (M1–M6)",
         "Architecture design, cloud infra setup, core API", "$2,100,000"],
        ["Phase 1: Foundation (M1–M6)",
         "IVA 3D base build (lobby + 3 zones)",              "$1,400,000"],
        ["Phase 2: Core Build (M7–M12)",
         "IVA full feature completion + AI integration",     "$2,800,000"],
        ["Phase 2: Core Build (M7–M12)",
         "Concierge 360° MVP (booking, smart home, chat)",   "$1,900,000"],
        ["Phase 3: Launch & Scale (M13–M18)",
         "AI fine-tuning, analytics, partner integrations",  "$1,300,000"],
        ["Phase 3: Launch & Scale (M13–M18)",
         "QA, security audit, load testing, go-live",        "$500,000"],
        ["Contingency (10%)",
         "Risk buffer across all phases",                    "$1,000,000"],
    ]
    make_table(doc, headers, rows, col_widths=[Cm(5), Cm(8), Cm(4)])

    # 7.5
    add_heading(doc, "7.5  Sensitivity Analysis", 2)
    headers = ["Scenario", "Conversion Rate", "Concierge MAR", "Year 3 Net Revenue", "3-Year ROI"]
    rows = [
        ["🐻 Bear Case",  "3.5% (vs 2% base)", "40% adoption", "$12,500,000", "+19%"],
        ["📊 Base Case",  "6.0% (vs 2% base)", "55% adoption", "$22,490,000", "+114%"],
        ["🐂 Bull Case",  "8.5% (vs 2% base)", "72% adoption", "$34,200,000", "+225%"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(3), Cm(3), Cm(3), Cm(4), Cm(3)])

    insert_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 – PHASED ROADMAP
# ══════════════════════════════════════════════════════════════════════════════

def add_section8(doc):
    add_heading(doc, "SECTION 8: PHASED ROADMAP", 1)
    headers = ["Phase", "Months", "Key Deliverables", "Budget", "KPI Targets"]
    rows = [
        ["Phase 1\nFoundation",
         "M1–M6",
         "• Cloud infra & DevOps\n• Core API gateway\n• IVA lobby + 3 zones\n• Design system",
         "$3,500,000",
         "• Dev environment live M2\n• First zone demo M5\n• Security audit pass M6"],
        ["Phase 2\nCore Build",
         "M7–M12",
         "• IVA full 6-zone completion\n• AI scoring engine v1\n• Concierge MVP\n• CRM integration",
         "$4,700,000",
         "• Soft launch IVA M10\n• Concierge beta M11\n• 100 test users M12"],
        ["Phase 3\nLaunch & Scale",
         "M13–M18",
         "• AI fine-tuning (VN NLP)\n• Full Concierge launch\n• Partner marketplace\n• Analytics v2",
         "$1,800,000",
         "• IVA Entry Rate ≥35% M14\n• Concierge MAR ≥40% M16\n• Full go-live M18"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(2.5), Cm(2), Cm(6), Cm(2.5), Cm(5)])

    add_heading(doc, "Key Milestones Timeline", 2)
    milestones = [
        ("M0",  "Project kickoff, vendor onboarding, architecture sign-off"),
        ("M2",  "Cloud infrastructure live, CI/CD pipelines operational"),
        ("M5",  "IVA Showroom Zone 1 demo ready for stakeholder review"),
        ("M8",  "AI Behavioral Scoring Engine v1 integrated into IVA"),
        ("M10", "IVA Soft Launch — internal sales team testing"),
        ("M11", "Concierge 360° Beta — 50 resident pilot group"),
        ("M14", "IVA Public Launch — marketing campaigns activated"),
        ("M16", "Concierge Full Launch — all HH4 residents onboarded"),
        ("M18", "Platform fully operational — all KPIs tracked, SOC active"),
    ]
    for ms, desc in milestones:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"[{ms}]  ")
        set_run_font(r1, 11, bold=True, color=GOLD)
        r2 = p.add_run(desc)
        set_run_font(r2, 11, color=BLACK)

    insert_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 – RFP & VENDOR SELECTION
# ══════════════════════════════════════════════════════════════════════════════

def add_section9(doc):
    add_heading(doc, "SECTION 9: RFP & VENDOR SELECTION", 1)

    # 9.1
    add_heading(doc, "9.1  Vendor Evaluation Scorecard", 2)
    headers = ["Criterion", "Weight", "Sub-criteria", "Scoring (1–5)", "Max Points"]
    rows = [
        ["Technical Capability",
         "30%",
         "AI/ML maturity, cloud-native, 3D/WebGL, mobile",
         "1–5 per sub", "30"],
        ["Vietnamese Market Experience",
         "20%",
         "Local projects, VN NLP, regulatory knowledge",
         "1–5 per sub", "20"],
        ["Financial Stability",
         "15%",
         "3-year audited P&L, bank guarantee capacity",
         "1–5 per sub", "15"],
        ["Delivery Track Record",
         "20%",
         "On-time delivery %, reference projects, case studies",
         "1–5 per sub", "20"],
        ["Total Cost of Ownership",
         "15%",
         "Licence model, support SLA, exit provisions",
         "1–5 per sub", "15"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(4), Cm(1.5), Cm(6), Cm(2.5), Cm(2)])

    # 9.2
    add_heading(doc, "9.2  Mandatory Requirements", 2)
    headers = ["Category", "Technical Requirements", "Commercial Requirements"]
    rows = [
        ["Security",
         "ISO 27001 or SOC 2 Type II certified",
         "Cyber liability insurance ≥ $5M"],
        ["Performance",
         "99.9% uptime SLA, <3s global load time",
         "Penalty clause: 5% fee per 0.1% breach"],
        ["AI/Data",
         "On-premise or VN data residency option",
         "IP ownership transferable to Splendora"],
        ["Integration",
         "REST/GraphQL APIs, Salesforce CRM connector",
         "6-month source code escrow"],
        ["Support",
         "24/7 NOC, <4h P1 response, Vietnamese-speaking",
         "3-year support contract minimum"],
    ]
    make_table(doc, headers, rows, col_widths=[Cm(3), Cm(7), Cm(7)])

    # 9.3
    add_heading(doc, "9.3  Contract Milestones & Payment Schedule", 2)
    headers = ["Milestone", "Month", "Deliverable", "Payment %"]
    rows = [
        ["M0 — Contract Sign",    "M0",  "Signed contract, project charter, team onboarding",   "15%"],
        ["M1 — Architecture",     "M1",  "Technical architecture approved, infra live",          "10%"],
        ["M2 — Design Approved",  "M3",  "UX/UI design system approved by Splendora",            "10%"],
        ["M3 — Alpha Build",      "M6",  "IVA Zone 1–3 functional, Concierge wireframes",        "15%"],
        ["M4 — Beta Release",     "M10", "Full IVA beta, Concierge MVP, AI v1 integrated",       "20%"],
        ["M5 — UAT Pass",         "M13", "UAT sign-off by Splendora QA team",                   "10%"],
        ["M6 — Go-Live",          "M18", "Public launch, all KPIs instrumented",                 "15%"],
        ["M7 — Stabilisation",    "M21", "90-day post-launch support, handover docs",            "5%"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(4), Cm(1.8), Cm(9), Cm(2)])

    insert_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 – RISK MITIGATION
# ══════════════════════════════════════════════════════════════════════════════

def add_section10(doc):
    add_heading(doc, "SECTION 10: RISK MITIGATION", 1)
    headers = ["Rủi ro", "Mức độ", "Giải pháp giảm thiểu", "Owner", "Timeline"]
    rows = [
        ["Vendor delivery delay",
         "🔴 High",
         "Milestone-based payments, 2nd vendor pre-qualified",
         "PMO",
         "Ongoing"],
        ["AI quality insufficient (VN NLP)",
         "🟡 Medium",
         "3-month fine-tuning buffer, fallback to GPT-4o API",
         "AI Lead",
         "M8–M12"],
        ["Low resident adoption of Concierge app",
         "🟡 Medium",
         "Gamification, onboarding concierge team, free trial",
         "Product",
         "M16+"],
        ["Data security breach",
         "🔴 High",
         "ISO 27001, penetration tests, cyber insurance",
         "CISO",
         "Ongoing"],
        ["Regulatory / PDPA non-compliance",
         "🟡 Medium",
         "DPO appointed, OneTrust platform, legal review",
         "Legal",
         "M1"],
        ["3D content quality below expectation",
         "🟡 Medium",
         "Milestone review at Zone 1, dedicated QA artist",
         "Creative Dir",
         "M5"],
        ["Integration failure with existing CRM/ERP",
         "🟡 Medium",
         "API contract testing, sandbox environment M2",
         "IT Arch",
         "M2–M6"],
        ["Budget overrun",
         "🔴 High",
         "10% contingency reserve, change control board",
         "CFO",
         "Ongoing"],
        ["Key staff attrition (vendor side)",
         "🟢 Low",
         "Key-person clause in contract, knowledge transfer",
         "PMO / HR",
         "Ongoing"],
        ["Market downturn affecting sales targets",
         "🟢 Low",
         "Platform ROI modelled at Bear Case scenario",
         "C-Suite",
         "Annual review"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(4), Cm(2), Cm(6), Cm(2.5), Cm(2)])

    insert_page_break(doc)


# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 – COMPETITIVE DIFFERENTIATION
# ══════════════════════════════════════════════════════════════════════════════

def add_section11(doc):
    add_heading(doc, "SECTION 11: COMPETITIVE DIFFERENTIATION", 1)
    add_body(doc,
        "The matrix below benchmarks 12 key capability dimensions of the HH4 Platform "
        "against two competitive proxies — Competitor A represents the leading domestic "
        "PropTech platform; Market Average represents the typical Vietnam luxury developer "
        "digital offering.")

    headers = ["Capability", "HH4 Platform", "Competitor A", "Market Average"]
    rows = [
        ["3D / Metaverse Showroom",
         "✅ Photorealistic WebGL, VR-ready",
         "⚠️ Basic 360° photos",
         "❌ Static images"],
        ["AI Virtual Guide",
         "✅ Vietnamese NLP, custom-trained",
         "⚠️ Rule-based chatbot",
         "❌ None"],
        ["Behavioral Lead Scoring",
         "✅ Real-time ML pipeline",
         "⚠️ Manual scoring",
         "❌ None"],
        ["Live Video Consultation",
         "✅ In-platform, recorded",
         "✅ 3rd-party Zoom link",
         "⚠️ Phone only"],
        ["Concierge / Resident App",
         "✅ Super-app with IoT",
         "⚠️ Basic maintenance app",
         "❌ None"],
        ["Smart Home Integration",
         "✅ Full IoT dashboard",
         "❌ Not available",
         "❌ None"],
        ["Loyalty & Rewards Engine",
         "✅ Points, tiers, partner redemption",
         "⚠️ Basic referral only",
         "❌ None"],
        ["Analytics & BI",
         "✅ Real-time, multi-source",
         "⚠️ Monthly reports",
         "⚠️ Basic CRM reports"],
        ["Multi-language Support",
         "✅ EN / VI / ZH",
         "⚠️ EN / VI only",
         "⚠️ VI only"],
        ["Data Sovereignty",
         "✅ VN data residency",
         "⚠️ SG-based",
         "❌ Not specified"],
        ["Security Certification",
         "✅ ISO 27001 + SOC 2 required",
         "⚠️ ISO 27001 only",
         "❌ None"],
        ["Open API / Ecosystem",
         "✅ GraphQL + REST, partner SDK",
         "⚠️ Limited webhooks",
         "❌ Closed system"],
    ]
    make_table(doc, headers, rows,
               col_widths=[Cm(4.5), Cm(5), Cm(4), Cm(4)])

    add_body(doc,
        "\nConclusion: The HH4 Platform is the only solution in the Vietnam market delivering "
        "a fully integrated, AI-first, metaverse-grade real estate experience combining "
        "pre-purchase journey optimisation with post-handover resident lifecycle management. "
        "This constitutes a durable competitive moat for Splendora's premium positioning.",
        italic=True)


# ══════════════════════════════════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════════════════════════════════

def main():
    doc = setup_document()
    add_header_footer(doc)

    add_cover(doc)
    add_toc(doc)
    add_section1(doc)
    insert_page_break(doc)
    add_section2(doc)
    add_section3(doc)
    add_section4(doc)
    add_section5(doc)
    add_section6(doc)
    add_section7(doc)
    add_section8(doc)
    add_section9(doc)
    add_section10(doc)
    add_section11(doc)

    # Output path relative to this script
    out_dir  = os.path.dirname(os.path.abspath(__file__))
    out_path = os.path.join(out_dir, "HH4_Platform_Solution_Architecture.docx")
    doc.save(out_path)
    print(f"✅  Document saved: {out_path}")


if __name__ == "__main__":
    main()
